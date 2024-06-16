import streamlit as st
from docx import Document
from PyPDF2 import PdfReader
# from gensim.summarization import summarize
import os
from transformers import pipeline
import base64
import DocumentGeneration, DocumentSummarization
def get_binary_file_downloader_html(content, filename):
    """
    Create a downloader HTML link.

    Args:
        content (str): Content to be downloaded.
        filename (str): Name of the file.

    Returns:
        str: HTML for downloader link.
    """
    b64 = base64.b64encode(content.encode()).decode()
    href = f'<a href="data:file/txt;base64,{b64}" download="{filename}">Download {filename}</a>'
    return href

def document_summarization(file):
    if file.name.endswith('.txt'):
        with open(file.name, 'r', encoding='utf-8') as f:
            text = f.read()
    elif file.name.endswith('.pdf'):
        pdf_reader = PdfReader(file)
        text = ''
        for page_num in range(len(pdf_reader.pages)):
            text += pdf_reader.pages[page_num].extract_text()
    elif file.name.endswith('.docx'):
        doc = Document(file)
        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append(para.text)
        text = '\n'.join(paragraphs)
    else:
        return "Unsupported file format. Please upload a .txt, .pdf, or .docx file."

    # summary = summarize(text)

    # summarizer = pipeline("summarization", model="t5-base")
    # summary = summarizer(text, max_length=1000, min_length=30, do_sample=False)
    # print(summary)
    summary = DocumentSummarization.Summarization(text)
    return summary

# def document_generation(document_type,scenario):
def document_generation(document_type, party_a, party_b, location, price, duration, scenario):
    requirements = {
    'document_type': document_type,
    'sections': [
        {
            'title': 'party_a',
            'content': party_a
        },
        {
            'title': 'party_b',
            'content': party_b
        },
        {
            'title': 'location',
            'content': location
        },
        {
            'title': 'price',
            'content': price
        },
        {
            'title': 'duration',
            'content': duration
        },
        {
            'title': 'scenario',
            'content': scenario
        }
    ]
}
    document = DocumentGeneration.Generation(requirements)
    print(document)
    return document

def display_document(document, doc_filename):
    st.markdown(f"## Displaying Document: {doc_filename}")
    # with open(document_path, 'rb') as file:
    #     document_content = file.read().decode(encoding, errors='replace')
    st.write(document)


st.title('AI Powered Legal Documentation Assistant')

option = st.sidebar.radio('Choose an option:', ['Home','Document Summarization', 'Document Generation'])

if option == 'Home':
    st.write("""
    This application is used for document generation and summarization of legal documents. 
    Please consider any misinterpretations as AI may make mistakes.
    
    For document generation, you can choose from 5 types:
    - Rental
    - Lease
    - MOU (Memorandum of Understanding)
    - Employee Contract
    - Loan
    """)

if option == 'Document Summarization':
    st.sidebar.header('Document Summarization')
    uploaded_file = st.sidebar.file_uploader('Upload Document', type=['txt', 'pdf', 'docx'])

    if uploaded_file is not None:
        if st.sidebar.button('Generate Summary'):
            summary = document_summarization(uploaded_file)
            st.subheader('Summary')
            st.write(summary)

elif option == 'Document Generation':
    st.sidebar.header('Document Generation')
    document_type = st.sidebar.selectbox('Document Type', ['Rental', 'Lease','MOU', 'Employee Contract', 'Loan'])
    party_a = st.sidebar.text_input('Parties Involved (A)')
    party_b = st.sidebar.text_input('Parties Involved (B)')
    location = st.sidebar.text_input('Location (If Required)')
    price = st.sidebar.text_input('Price (If Required)')
    duration = st.sidebar.text_input('Duration (If Required)')
    scenario = st.sidebar.text_area('Scenario')

    if st.sidebar.button('Generate Document'):
        # document_path = document_generation(document_type, scenario)
        document, doc_filename = document_generation(document_type, party_a, party_b, location, price, duration, scenario)

        st.write(doc_filename)
        st.subheader('Generated Document')
        display_document(document, doc_filename)
        st.sidebar.markdown(get_binary_file_downloader_html(doc_filename, 'Document'), unsafe_allow_html=True)


# document_type = "Agreement"
# party_a = "John Doe"
# party_b = "Jane Smith"
# location = "New York"
# price = "1000"
# duration = "6 months"
# scenario = "This agreement is for the lease of a commercial property located in downtown New York. The agreement is between John Doe, represented by ABC Corporation, and Jane Smith, represented by XYZ Enterprises. The lease term is for 6 months at a monthly rent of $1000. The property is located at 123 Main Street, New York."


from transformers import BartForConditionalGeneration, BartTokenizer
from docx import Document

def generate_legal_points(prompt):
    # Initialize BART model and tokenizer
    tokenizer = BartTokenizer.from_pretrained('facebook/bart-large-cnn')
    model = BartForConditionalGeneration.from_pretrained('facebook/bart-large-cnn')

    legal_points = []

    # Generate legal points based on the prompt
    input_ids = tokenizer.encode(prompt, return_tensors="pt", max_length=512, truncation=True)
    
    # Generate 20 unique sentences
    output = model.generate(input_ids, max_length=100, num_return_sequences=10, num_beams=10, do_sample=True, temperature=1.0, top_p=0.95)

    # Decode the generated legal points
    for i in range(len(output)):
        point = tokenizer.decode(output[i], skip_special_tokens=True)
        if point not in legal_points:
            legal_points.append(point)
            if len(legal_points) >= 20:
                break

    return legal_points

def generate_document(document_type, party_a, party_b, location, price, duration, scenario):
    result = ''
    # Generate the scenario with placeholders replaced by input values
    scenario = scenario.format(document_type=document_type,party_a=party_a, party_b=party_b, duration=duration, price=price, location=location)

    # Generate legal points based on the scenario
    legal_points = generate_legal_points(scenario)
    
    # Create a new Document
    doc = Document()
    doc.add_heading(f"{document_type} Agreement", level=1)
    result += f"{document_type} Agreement"
    result += '\n'
    result += "\n\n\n"
    # Add scenario
    doc.add_paragraph(scenario)
    result += scenario
    result += '\n'
    result += "\n\n\n"
    result += "\n\n\n"
    # Add generated legal points
    for point in legal_points:
        doc.add_paragraph(point)
        result += point
        result += '\n'
        result += "\n\n\n"
        result += "\n\n\n"
    # Add signatures and addresses
    doc.add_paragraph("\n\n\n")
    result += "\n\n\n"
    result += "\n\n\n"
    doc.add_paragraph("Signed and agreed by:")
    result += "Signed and agreed by:\n\n"
    result += "\n\n\n"
    result += "\n\n\n"
    doc.add_paragraph("\n\n")
    result += f"{party_a}: ________________________\n"
    result += "\n\n\n"
    result += "\n\n\n"
    doc.add_paragraph(f"{party_a}: ________________________\n")
    result += "Authorized Signature: ________________________\n"
    result += "\n\n\n"
    result += "\n\n\n"
    doc.add_paragraph("Authorized Signature: ________________________\n")
    result += "Date: ________________________\n"
    result += "\n\n\n"
    result += "\n\n\n"
    doc.add_paragraph("Date: ________________________\n")
    result += "Address: ________________________\n\n\n"
    result += "\n\n\n"
    result += "\n\n\n"
    doc.add_paragraph("Address: ________________________\n\n\n")
    doc.add_paragraph(f"{party_b}: ________________________\n")
    result += f"{party_b}: ________________________\n"
    result += "\n\n\n"
    result += "\n\n\n"
    doc.add_paragraph("Authorized Signature: ________________________\n")
    result += "Authorized Signature: ________________________\n"
    result += "\n\n\n"
    result += "\n\n\n"
    doc.add_paragraph("Date: ________________________\n")
    result += "Date: ________________________\n"
    result += "\n\n\n"
    result += "\n\n\n"
    result += "Address: ________________________"
    result += "\n\n\n"
    result += "\n\n\n"
    doc.add_paragraph("Address: ________________________")

    # Save the document
    doc_filename = f"{document_type}_document.docx"
    doc.save(doc_filename)

    return result, doc_filename

def assign_values(requirements):
    params = {}
    for section in requirements['sections']:
        title = section['title']
        content = section['content']
        # if title == 'document_type':
        #     params['document_type'] = content
        if title == 'party_a':
            params['party_a'] = content
        elif title == 'party_b':
            params['party_b'] = content
        elif title == 'location':
            params['location'] = content
        elif title == 'price':
            params['price'] = content
        elif title == 'duration':
            params['duration'] = content
        elif title == 'scenario':
            params['scenario'] = content
    return params

def Generation(requirements):

    params = assign_values(requirements)

    document_type = requirements['document_type']
    party_a = params['party_a']
    party_b = params['party_b']
    location = params['location']
    price = params['price']
    duration = params['duration']
    scenario = params['scenario']

    generated_document = generate_document(document_type, party_a, party_b, location, price, duration, scenario)

    print("Document generated:", generated_document)

    return generated_document
