from transformers import BartForConditionalGeneration, BartTokenizer
from docx import Document
import random

def generate_legal_points(prompts):
    # Initialize BART model and tokenizer
    tokenizer = BartTokenizer.from_pretrained('facebook/bart-large-cnn')
    model = BartForConditionalGeneration.from_pretrained('facebook/bart-large-cnn')

    legal_points = []

    for prompt in prompts:
        # Generate legal points based on the prompt
        input_ids = tokenizer.encode(prompt, return_tensors="pt", max_length=512, truncation=True)
        
        # Generate multiple sentences to ensure variety
        output = model.generate(input_ids, max_length=100, num_return_sequences=1, num_beams=1, do_sample=True, temperature=1.5, top_p=0.9)
        # print(output)
        # Decode the generated legal points and ensure uniqueness
        for o in output:
            point = tokenizer.decode(o, skip_special_tokens=True)
            if point not in legal_points:
                legal_points.append(point)
                # if len(legal_points) >= 20:
                #     break

        # if len(legal_points) >= 20:
        #     break

    return legal_points

def generate_document(document_type, party_a, party_b, location, price, duration, scenario):
    result = ''
    # Create different prompts based on input fields
    prompts = [
        f"Describe the legal obligations of {party_a} and {party_b} under a {document_type} agreement in {location}.",
        f"Explain the terms and conditions of the {document_type} agreement between {party_a} and {party_b} for {price}.",
        f"Highlight the legal responsibilities of {party_b} in the {document_type} agreement lasting {duration} at {location}.",
        f"Detail the compliance requirements for {party_a} in a {document_type} agreement valued at {price} for {duration}.",
        f"Summarize the key legal points of the {document_type} agreement between {party_a} and {party_b} at {location}.",
        f"Generate a legal point based on this scenario: {scenario}",
        f"Outline potential legal conflicts in the {document_type} agreement involving {party_a}, {party_b}, and a property in {location}.",
        f"Analyze a legal situation for {party_a} and {party_b} in a {document_type} agreement over {duration} at {location}."
    ]

    # Shuffle prompts to add more variability
    random.shuffle(prompts)

    # Generate legal points based on the prompts
    legal_points = generate_legal_points(prompts)
    
    # Create a new Document
    doc = Document()
    doc.add_heading(f"{document_type} Agreement", level=1)
    result += f"{document_type} Agreement\n\n\n"
    
    # Add scenario
    doc.add_paragraph(scenario)
    result += scenario + '\n\n\n\n\n\n'
    
    # Add generated legal points
    for point in legal_points[:4]:  # Add only the first 4 unique points to reduce the length
        doc.add_paragraph(point)
        result += point + '\n\n\n\n\n\n'
    
    # Add signatures and addresses
    doc.add_paragraph("\n\n\n")
    result += "\n\n\n\n\n\n"
    doc.add_paragraph("Signed and agreed by:")
    result += "Signed and agreed by:\n\n\n\n\n\n"
    doc.add_paragraph("\n\n")
    result += f"{party_a}: ________________________\n\n\n\n\n\n"
    doc.add_paragraph(f"{party_a}: ________________________\n")
    result += "Authorized Signature: ________________________\n\n\n\n\n\n"
    doc.add_paragraph("Authorized Signature: ________________________\n")
    result += "Date: ________________________\n\n\n\n\n\n"
    doc.add_paragraph("Date: ________________________\n")
    result += "Address: ________________________\n\n\n\n\n\n"
    doc.add_paragraph("Address: ________________________\n\n\n")
    doc.add_paragraph(f"{party_b}: ________________________\n")
    result += f"{party_b}: ________________________\n\n\n\n\n\n"
    doc.add_paragraph("Authorized Signature: ________________________\n")
    result += "Authorized Signature: ________________________\n\n\n\n\n\n"
    doc.add_paragraph("Date: ________________________\n")
    result += "Date: ________________________\n\n\n\n\n\n"
    result += "Address: ________________________\n\n\n\n\n\n"
    doc.add_paragraph("Address: ________________________")

    # Save the document
    doc_filename = f"{document_type}_document.docx"
    doc.save(doc_filename)

    return result, doc_filename

def assign_values(requirements):
    params = {}
    for section in requirements['sections']:
        title = section['title']
        content = section['content']
        if title == 'party_a':
            params['party_a'] = content
        elif title == 'party_b':
            params['party_b'] = content
        elif title == 'location':
            params['location'] = content
        elif title == 'price':
            params['price'] = content
        elif title == 'duration':
            params['duration'] = content
        elif title == 'scenario':
            params['scenario'] = content
    return params

def Generation(requirements):

    params = assign_values(requirements)

    document_type = requirements['document_type']
    party_a = params['party_a']
    party_b = params['party_b']
    location = params['location']
    price = params['price']
    duration = params['duration']
    scenario = params['scenario']

    generated_document, doc_filename = generate_document(document_type, party_a, party_b, location, price, duration, scenario)

    print("Document generated:", generated_document)

    return generated_document, doc_filename
